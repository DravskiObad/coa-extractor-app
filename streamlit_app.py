import streamlit as st
import pandas as pd
import os
import tempfile
import uuid
from extractor import process_single_file
from docx import Document
from io import BytesIO
import time

# Page configuration
st.set_page_config(
    page_title="CoA Extraction App",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #455e6f;
        text-align: center;
        margin-bottom: 2rem;
        background: linear-gradient(135deg, #455e6f 0%, #38b292 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .instructions-box {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border: 2px solid #38b292;
        border-radius: 15px;
        padding: 20px;
        margin-bottom: 20px;
    }
    
    .stAlert {
        border-radius: 10px;
    }
    
    .footer {
        text-align: center;
        color: #455e6f;
        font-size: 0.9rem;
        margin-top: 40px;
        padding: 20px;
        border-top: 1px solid rgba(69, 94, 111, 0.1);
    }
</style>
""", unsafe_allow_html=True)

def allowed_file(filename):
    """Check if file extension is allowed"""
    return filename.lower().endswith('.pdf')

def restructure_data(all_extracted_data):
    """Restructure extracted data for display"""
    parameters = {}
    batch_numbers = set()
    
    for file_data in all_extracted_data:
        if not file_data or 'data' not in file_data:
            continue
            
        extracted_items = file_data['data']
        
        if not extracted_items:
            continue
            
        for item in extracted_items:
            if not item or not isinstance(item, dict):
                continue
                
            if 'batch_number' in item and 'analytical_parameters' in item:
                batch_number = item.get('batch_number', '')
                if batch_number:
                    batch_numbers.add(batch_number)
                
                analytical_params = item.get('analytical_parameters', [])
                if isinstance(analytical_params, list):
                    for param in analytical_params:
                        if isinstance(param, dict):
                            param_name = param.get('parameter', '')
                            result = param.get('result', '')
                            method = param.get('method', '')
                            
                            if param_name:
                                if param_name not in parameters:
                                    parameters[param_name] = {
                                        'method': method,
                                        'batches': {}
                                    }
                                
                                if batch_number:
                                    parameters[param_name]['batches'][batch_number] = result
    
    # Convert to table format
    batch_numbers = sorted(list(batch_numbers))
    table_data = []
    
    for param, data in parameters.items():
        row = {'Parameter (unit)': param}
        for batch in batch_numbers:
            row[batch] = data['batches'].get(batch, '')
        row['Method of analysis'] = data['method']
        table_data.append(row)
    
    return table_data, batch_numbers

def create_word_document(table_data, batch_numbers):
    """Generate Word document from table data"""
    doc = Document()
    doc.add_heading('Certificate of Analysis Data', 0)
    
    if batch_numbers:
        headers = ['Parameter (unit)'] + batch_numbers + ['Method of analysis']
        # Create table with 2 header rows
        table = doc.add_table(rows=2, cols=len(headers))
        table.style = 'Table Grid'
        
        # First header row - merge cells for title
        first_row = table.rows[0].cells
        first_row[0].text = 'Parameter (unit)'
        
        # Merge batch number cells and add title
        if len(batch_numbers) > 1:
            first_row[1].text = 'Batch number'
            for i in range(2, len(batch_numbers) + 1):
                first_row[1].merge(first_row[i])
        else:
            first_row[1].text = 'Batch number'
        
        first_row[-1].text = 'Method of analysis'
        
        # Second header row - individual batch numbers
        second_row = table.rows[1].cells
        second_row[0].text = ''  # Empty for parameter column
        for i, batch in enumerate(batch_numbers):
            second_row[i + 1].text = batch
        second_row[-1].text = ''  # Empty for method column
    else:
        headers = ['Parameter (unit)', 'Method of analysis']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
    
    # Add data rows
    for row_data in table_data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            row_cells[i].text = str(row_data.get(header, ''))
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

# Initialize session state
if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = None
if 'table_data' not in st.session_state:
    st.session_state.table_data = None
if 'batch_numbers' not in st.session_state:
    st.session_state.batch_numbers = None

# Main app
st.markdown('<h1 class="main-header">CoA Extraction App</h1>', unsafe_allow_html=True)

# Instructions
st.markdown("""
<div class="instructions-box">
    <h3>📋 Instructions</h3>
    <ul>
        <li><strong>⚠️ Please split CoAs into smaller files, up to 10 pages long at max.</strong></li>
        <li>📄 Upload PDF files containing Certificates of Analysis (CoAs)</li>
        <li>📁 Supported formats: PDF files only</li>
        <li>📤 Multiple files can be uploaded simultaneously</li>
        <li>🤖 The AI will automatically extract analytical parameters, batch numbers, and methods</li>
        <li>📊 Results will be presented in a standardized table format</li>
        <li>💾 Download the extracted data as a Word document for further use</li>
    </ul>
</div>
""", unsafe_allow_html=True)

# File upload section
st.subheader("📤 Upload PDF Files")
uploaded_files = st.file_uploader(
    "Choose PDF files",
    type="pdf",
    accept_multiple_files=True,
    help="Select one or more PDF files containing Certificate of Analysis documents"
)

# Process files
if uploaded_files:
    if st.button("🚀 Extract Data", type="primary"):
        # Validate files
        valid_files = []
        for file in uploaded_files:
            if allowed_file(file.name):
                valid_files.append(file)
            else:
                st.error(f"❌ File type not allowed: {file.name}. Only PDF files are supported.")
        
        if valid_files:
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            all_extracted_data = []
            total_files = len(valid_files)
            
            for i, file in enumerate(valid_files):
                status_text.text(f"Processing {file.name}... ({i+1}/{total_files})")
                progress_bar.progress((i) / total_files)
                
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(file.getvalue())
                    tmp_file_path = tmp_file.name
                
                try:
                    # Process the file
                    extracted_data = process_single_file(tmp_file_path)
                    
                    if extracted_data:
                        all_extracted_data.append({
                            'filename': file.name,
                            'data': extracted_data
                        })
                        st.success(f"✅ Successfully processed: {file.name}")
                    else:
                        st.warning(f"⚠️ No data extracted from: {file.name}")
                        
                except Exception as e:
                    st.error(f"❌ Error processing {file.name}: {str(e)}")
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
            
            progress_bar.progress(1.0)
            status_text.text("Processing complete!")
            
            if all_extracted_data:
                # Restructure data for display
                table_data, batch_numbers = restructure_data(all_extracted_data)
                
                # Store in session state
                st.session_state.extracted_data = all_extracted_data
                st.session_state.table_data = table_data
                st.session_state.batch_numbers = batch_numbers
                
                st.success(f"🎉 Successfully processed {len(all_extracted_data)} files!")
                time.sleep(1)
                st.rerun()
            else:
                st.error("❌ No data could be extracted from any of the uploaded files.")

# Display results
if st.session_state.table_data:
    st.subheader("📊 Extracted Data")
    
    # Convert to DataFrame for better display
    df = pd.DataFrame(st.session_state.table_data)
    
    # Reorder columns to match the original layout
    if st.session_state.batch_numbers:
        column_order = ['Parameter (unit)'] + st.session_state.batch_numbers + ['Method of analysis']
        df = df[column_order]
    
    # Display the table
    st.dataframe(
        df,
        use_container_width=True,
        hide_index=True
    )
    
    # Download button
    if st.button("💾 Download as Word Document", type="secondary"):
        try:
            doc_io = create_word_document(st.session_state.table_data, st.session_state.batch_numbers)
            
            st.download_button(
                label="📥 Click to Download",
                data=doc_io.getvalue(),
                file_name="coa_data.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("✅ Document ready for download!")
        except Exception as e:
            st.error(f"❌ Error creating document: {str(e)}")

# Clear data button
if st.session_state.table_data:
    if st.button("🗑️ Clear Data and Start New Upload"):
        st.session_state.extracted_data = None
        st.session_state.table_data = None
        st.session_state.batch_numbers = None
        st.rerun()

# Footer
st.markdown("""
<div class="footer">
    Powered by European Food Safety Authority 2025
</div>
""", unsafe_allow_html=True)